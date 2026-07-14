import io
from fastapi import UploadFile
from langchain_core.documents import Document
from pypdf import PdfReader
from langchain_text_splitters import RecursiveCharacterTextSplitter
from docx import Document as DocxReader
from typing import List

async def file_input(file: UploadFile) -> List[Document]:

    try:
        filename = file.filename
        file_bytes = await file.read()

        if filename.lower().strip().endswith('.pdf'):
            pdf = PdfReader(io.BytesIO(file_bytes))
            pdf_file = []

            for page in pdf.pages:
                content = page.extract_text()
                if content:
                    pdf_file.append(content)

            pdf_file = "\n".join(pdf_file)
            doc = Document(
                page_content = pdf_file,
                metadata = {'filename': filename}
            )
        
        if filename.lower().strip().endswith('.txt'):
            content = file_bytes.decode("utf-8", errors="ignore")

            doc = Document(page_content= content, metadata= {'filename': filename})
        
        if filename.lower().strip().endswith('.docx'):
            docx = DocxReader(io.BytesIO(file_bytes))
            docx_file = []

            for para in docx.paragraphs:
                content = para.text
                if content:
                    docx_file.append(content)
                    
            docx_file = '\n'.join(docx_file)
            doc = Document(
                page_content = docx_file,
                metadata = {'filename': filename}
            )

        text_content = doc.page_content
        docs = [doc]

        if len(text_content) < 4000:
            return docs
        
        splitter = RecursiveCharacterTextSplitter(
            chunk_size = 800,
            chunk_overlap = 300
        )
        
        final_doc = splitter.split_documents(docs)
        return final_doc
    
    except Exception as error:
        raise Exception(f"Failed to parse file: {str(error)}")